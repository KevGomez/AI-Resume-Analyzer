import os
import logging
import fitz  # PyMuPDF
from docx import Document
from werkzeug.utils import secure_filename
from typing import Optional, Tuple

# Configure logging
logger = logging.getLogger(__name__)

class FileService:
    def __init__(self, upload_folder: str):
        self.upload_folder = upload_folder
        if not os.path.exists(upload_folder):
            os.makedirs(upload_folder)
            logger.info(f"Created upload folder: {upload_folder}")

    def save_file(self, file) -> Tuple[bool, str]:
        """Save uploaded file and return success status and file_name"""
        try:
            if not file:
                logger.error("No file provided")
                return False, "No file provided"

            file_name = secure_filename(str(file.filename))
            if not file_name:
                logger.error("Invalid file name")
                return False, "Invalid file name"

            filepath = os.path.join(self.upload_folder, file_name)
            file.save(filepath)
            logger.info(f"Successfully saved file: {filepath}")
            return True, file_name
        except Exception as e:
            logger.error(f"Error saving file: {str(e)}")
            return False, str(e)

    def extract_text_from_pdf(self, filepath: str) -> Optional[str]:
        """Extract text from PDF file"""
        try:
            logger.info(f"Extracting text from PDF: {filepath}")
            pdf_document = fitz.Document(filepath)
            text = ""
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                # PyMuPDF 1.23.8+ uses get_text()
                text += page.get_text()  # type: ignore
            pdf_document.close()
            
            if not text.strip():
                logger.warning("Extracted empty text from PDF")
                return None
                
            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            return None

    def extract_text_from_docx(self, filepath: str) -> Optional[str]:
        """Extract text from DOCX file"""
        try:
            logger.info(f"Extracting text from DOCX: {filepath}")
            doc = Document(filepath)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            if not text.strip():
                logger.warning("Extracted empty text from DOCX")
                return None
                
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            return None

    def extract_text(self, file_name: str) -> Optional[str]:
        """Extract text from uploaded file"""
        try:
            logger.info(f"Starting text extraction for file: {file_name}")
            filepath = os.path.join(self.upload_folder, file_name)
            
            if not os.path.exists(filepath):
                logger.error(f"File not found: {filepath}")
                return None

            file_extension = os.path.splitext(file_name)[1].lower()
            
            if file_extension == '.pdf':
                return self.extract_text_from_pdf(filepath)
            elif file_extension == '.docx':
                return self.extract_text_from_docx(filepath)
            else:
                logger.error(f"Unsupported file extension: {file_extension}")
                return None
        except Exception as e:
            logger.error(f"Error in extract_text: {str(e)}")
            return None 